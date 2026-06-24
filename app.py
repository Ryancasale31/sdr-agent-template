"""
WBR SDR Agent — multi-event sponsorship pipeline
Run with: streamlit run app.py
"""
import streamlit as st
import json
import csv
import os
import importlib
from pathlib import Path
from dotenv import load_dotenv
import anthropic
from tavily import TavilyClient
import outlook_integration as outlook
import storage
from events_registry import EVENTS

load_dotenv()

# ── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="WBR SDR Agent",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── WBR brand theme ───────────────────────────────────────────────────────────
WBR_NAVY = "#0E2747"
WBR_BLUE = "#1E5BA8"
WBR_ORANGE = "#F39200"

def inject_theme():
    st.markdown(f"""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700;800&display=swap');

    html, body, [class*="css"] {{ font-family: 'Inter', sans-serif; }}

    /* Sidebar: WBR navy gradient with light text */
    section[data-testid="stSidebar"] > div {{
        background: linear-gradient(180deg, {WBR_NAVY} 0%, #163B6E 60%, {WBR_BLUE} 130%);
    }}
    section[data-testid="stSidebar"] h1,
    section[data-testid="stSidebar"] h2,
    section[data-testid="stSidebar"] h3,
    section[data-testid="stSidebar"] p,
    section[data-testid="stSidebar"] label,
    section[data-testid="stSidebar"] span,
    section[data-testid="stSidebar"] [data-testid="stMetricValue"],
    section[data-testid="stSidebar"] [data-testid="stMetricLabel"] {{
        color: #EAF1FB !important;
    }}
    /* keep dropdown/inputs readable (dark text on their own white box) */
    section[data-testid="stSidebar"] [data-baseweb="select"] *,
    section[data-testid="stSidebar"] input {{ color: #0E2747 !important; }}

    /* WBR wordmark in the sidebar */
    .wbr-logo {{ font-weight: 800; font-size: 2.1rem; letter-spacing: 1px; color: #fff; line-height: 1; }}
    .wbr-logo .dot {{ color: {WBR_ORANGE}; }}
    .wbr-sub {{ font-size: .68rem; letter-spacing: 2px; color: #9FB8DA; text-transform: uppercase; margin-top: 2px; }}

    /* Hero banner on the main area */
    .hero {{
        border-radius: 14px; padding: 26px 30px; margin-bottom: 14px; color: #fff;
        background: linear-gradient(110deg, {WBR_NAVY} 0%, {WBR_BLUE} 70%), url('app/static/hero.jpg');
        background-size: cover; background-position: center; background-blend-mode: multiply;
        box-shadow: 0 6px 20px rgba(14,39,71,.18);
    }}
    .hero h1 {{ margin: 0; font-size: 1.7rem; font-weight: 800; }}
    .hero p {{ margin: 4px 0 0; color: #CFE0F5; font-size: .95rem; }}
    .hero .pill {{ display:inline-block; background:{WBR_ORANGE}; color:#fff; font-weight:700;
        font-size:.72rem; padding:3px 10px; border-radius:20px; margin-top:10px; }}

    /* Primary buttons in WBR orange */
    .stButton > button[kind="primary"] {{ background:{WBR_ORANGE}; border-color:{WBR_ORANGE}; }}
    .stButton > button[kind="primary"]:hover {{ background:#d97f00; border-color:#d97f00; }}

    /* Tabs accent */
    .stTabs [data-baseweb="tab-highlight"] {{ background-color: {WBR_ORANGE}; }}
    .stTabs [aria-selected="true"] {{ color: {WBR_NAVY} !important; font-weight: 700; }}
    </style>
    """, unsafe_allow_html=True)

inject_theme()

# ── Event selector / login ────────────────────────────────────────────────────
def _show_login():
    st.markdown("""
    <style>
    .login-wrap { max-width: 420px; margin: 80px auto; padding: 36px 40px;
                  background: white; border-radius: 16px;
                  box-shadow: 0 8px 32px rgba(14,39,71,.12); }
    .login-title { font-size: 1.6rem; font-weight: 800; color: #0E2747; margin-bottom: 4px; }
    .login-sub   { font-size: .88rem; color: #6B7280; margin-bottom: 24px; }
    </style>
    <div class="login-wrap">
      <div class="login-title">WBR SDR Agent</div>
      <div class="login-sub">Select your event and enter the access code.</div>
    </div>
    """, unsafe_allow_html=True)

    col_l, col_c, col_r = st.columns([1, 2, 1])
    with col_c:
        st.markdown("#### Select Event")
        selected = st.selectbox(
            "Event",
            options=list(EVENTS.keys()),
            format_func=lambda k: EVENTS[k]["name"],
            label_visibility="collapsed",
        )
        password = st.text_input("Access code", type="password", placeholder="Access code")
        if st.button("Enter", type="primary", use_container_width=True):
            try:
                expected = st.secrets.get("event_passwords", {}).get(selected, "")
            except Exception:
                expected = ""
            # Also allow empty password in local dev (no secrets configured)
            if password == expected or (not expected and not password):
                st.session_state["event_id"] = selected
                st.session_state["event_cfg"] = EVENTS[selected]
                st.rerun()
            else:
                st.error("Incorrect access code.")

if "event_id" not in st.session_state:
    _show_login()
    st.stop()

# ── Active event ──────────────────────────────────────────────────────────────
_event_id  = st.session_state["event_id"]
_event_cfg = st.session_state["event_cfg"]

# ── Data files ────────────────────────────────────────────────────────────────
ICP_FILE = "icp_summary.json"
PIPELINE_FILE = "pipeline.json"

# ── Load ICP ──────────────────────────────────────────────────────────────────
def load_icp():
    return storage.load_icp(event_id=_event_id)

# ── Pipeline helpers ──────────────────────────────────────────────────────────
def load_json(path, default):
    p = Path(path)
    if not p.exists():
        return default
    with open(p) as f:
        return json.load(f)

def load_pipeline():
    return storage.load_pipeline(event_id=_event_id)

def save_pipeline(pipeline):
    storage.save_pipeline(pipeline, event_id=_event_id)

def get_company(pipeline, name):
    return next((c for c in pipeline if c["company"].lower() == name.lower()), None)

def upsert_company(pipeline, company_data):
    for i, c in enumerate(pipeline):
        if c["company"].lower() == company_data["company"].lower():
            pipeline[i] = {**c, **company_data}
            return pipeline
    pipeline.append(company_data)
    return pipeline

# ── AI helpers ────────────────────────────────────────────────────────────────
def research_company(company_name: str, icp: dict) -> dict:
    tavily = TavilyClient(api_key=os.getenv("TAVILY_API_KEY"))
    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    ecfg = st.session_state.get("event_cfg", {})

    # Web research
    search_kw = ecfg.get("search_keywords", "software products customers target market")
    results = tavily.search(
        query=f"{company_name} {search_kw}",
        max_results=5,
        search_depth="advanced",
    )
    web_context = "\n\n".join([
        f"Source: {r['url']}\n{r['content']}"
        for r in results.get("results", [])
    ])

    event_label = f"{ecfg.get('name','the event')} ({ecfg.get('location','')}, {ecfg.get('dates','')})"
    prompt = f"""You are a sponsorship sales analyst for {event_label}.

EVENT BUYER PROFILE:
- {icp['buyer_count']} registered buyers, {icp['senior_buyer_pct']}% VP/Director/SVP level
- Industries: Industrial/Manufacturing Equipment (41), Medical/Life Sciences (30), Tech/IT (17), Utilities (13)
- Top companies attending: {', '.join(icp['top_companies'][:10])}
- Top titles: {', '.join(icp['top_titles'][:8])}
- Existing sponsors: {', '.join(set(icp['existing_sponsors']))}

WEB RESEARCH ON {company_name}:
{web_context}

Based on this research, analyze {company_name} as a potential sponsor.

Return ONLY valid JSON:
{{
  "company": "{company_name}",
  "what_they_do": "<1-2 sentence plain English description>",
  "who_they_sell_to": "<their target buyer persona>",
  "score": <0-100>,
  "tier": "<A|B|C>",
  "fit_reason": "<2-3 sentences on why they fit this audience>",
  "pitch_angle": "<the ONE strongest reason they should sponsor -- specific to our attendee list>",
  "risk": "<one sentence on potential objection>",
  "status": "researched"
}}
"""

    message = client.messages.create(
        model="claude-opus-4-5",
        max_tokens=800,
        messages=[{"role": "user", "content": prompt}],
    )

    raw = message.content[0].text.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    return json.loads(raw.strip())


def generate_meeting_email(company_data: dict, contact_name: str, contact_title: str, icp: dict) -> dict:
    """Generate a single, direct email asking for a meeting -- uses the pitch angle as the hook."""
    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    ecfg = st.session_state.get("event_cfg", {})

    pitch_angle = company_data.get("pitch_angle") or company_data.get("outreach_note") or company_data.get("fit_reason") or ""
    vs_sponsor = company_data.get("vs_sponsor", "")
    competitor_line = f"Their competitors already sponsoring: {vs_sponsor}." if vs_sponsor else ""

    sender = ecfg.get("sender_name", "Your Name")
    event_label = f"{ecfg.get('name','the event')} ({ecfg.get('location','')}, {ecfg.get('dates','')})"
    prompt = f"""You are {sender}, a sponsorship sales rep for {event_label}.

EVENT AUDIENCE:
- {icp['buyer_count']} registered buyers, {icp['senior_buyer_pct']}% VP/Director/SVP level
- Top attending companies: {', '.join(icp['top_companies'][:10])}
- Top titles: {', '.join(icp['top_titles'][:8])}
- Already sponsoring: {', '.join(set(icp['existing_sponsors']))}

TARGET:
- Company: {company_data['company']}
- Contact: {contact_name}, {contact_title}
- What they do: {company_data.get('what_they_do', '')}
- Who they sell to: {company_data.get('who_they_sell_to', '')}
- Why they fit: {pitch_angle}
{competitor_line}

Write ONE short, direct email asking for a 15-minute call to explore sponsorship.
- Lead with the ONE most relevant reason their buyers are in the room (from the pitch angle above)
- Be specific -- name actual attendee titles or companies if it strengthens the case
- Ask for a specific, low-friction action: "15 minutes this week or next?"
- Under 120 words total
- Tone: peer-to-peer, confident, no fluff
- Never use: "I hope this email finds you well", "synergy", "leverage", "cutting-edge", "robust"
- Sign off: {sender} | {ecfg.get('event_brand', ecfg.get('name',''))}

Return ONLY valid JSON:
{{"subject": "...", "body": "..."}}
"""

    message = client.messages.create(
        model="claude-opus-4-5",
        max_tokens=600,
        messages=[{"role": "user", "content": prompt}],
    )

    raw = message.content[0].text.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    return json.loads(raw.strip())


def generate_emails(company_data: dict, contact_name: str, contact_title: str, icp: dict) -> list:
    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    ecfg = st.session_state.get("event_cfg", {})

    # Build competitor context if available
    vs_sponsor = company_data.get("vs_sponsor", "")
    competitor_line = f"- Their competitors already sponsoring this event: {vs_sponsor}" if vs_sponsor else ""

    # Outreach note from the prospect list
    outreach_note = company_data.get("outreach_note", "")
    outreach_line = f"- Internal outreach note: {outreach_note}" if outreach_note else ""

    # Category context
    category = company_data.get("category", "")
    category_line = f"- Their market category: {category}" if category else ""

    sender = ecfg.get("sender_name", "Your Name")
    event_label = f"{ecfg.get('name','the event')} ({ecfg.get('location','')}, {ecfg.get('dates','')})"
    event_brand = ecfg.get("event_brand", ecfg.get("name", "the event"))

    prompt = f"""You are a senior sponsorship sales rep for {event_label}.

EVENT AUDIENCE:
- {icp['buyer_count']} registered buyers, {icp['senior_buyer_pct']}% VP/Director/SVP level
- Top attending companies: {', '.join(icp['top_companies'][:12])}
- Top titles: {', '.join(icp['top_titles'][:10])}
- Already sponsoring: {', '.join(set(icp['existing_sponsors']))}

TARGET COMPANY INTEL:
- Name: {company_data['company']}
- What they do: {company_data.get('what_they_do', '')}
- Who they sell to: {company_data.get('who_they_sell_to', company_data.get('what_they_do', ''))}
- Best pitch angle: {company_data.get('pitch_angle', company_data.get('outreach_note', ''))}
{competitor_line}
{outreach_line}
{category_line}

CONTACT: {contact_name}, {contact_title}

Write a 3-touch cold email sequence to sell them a sponsorship.

Rules:
- Email 1: Lead with ONE specific audience insight most relevant to them. Under 150 words. Curiosity only, no hard pitch.
- Email 2 (Day 4): Connect their product to specific buyer titles/companies attending. If their competitors are sponsoring, mention it as social proof -- not as a threat. Under 175 words.
- Email 3 (Day 9): Soft close, limited spots, reference existing sponsors as validation. Under 100 words.
- Tone: Direct, peer-to-peer, confident. No fluff, no corporate speak.
- Never use: "I hope this email finds you well", "synergy", "leverage", "cutting-edge", "robust", "game-changing"
- Always sign off: {sender}, {event_brand}

Return ONLY valid JSON:
[
  {{"touch": 1, "send_day": "Day 1", "subject": "...", "body": "..."}},
  {{"touch": 2, "send_day": "Day 4", "subject": "...", "body": "..."}},
  {{"touch": 3, "send_day": "Day 9", "subject": "...", "body": "..."}}
]
"""

    message = client.messages.create(
        model="claude-opus-4-5",
        max_tokens=1500,
        messages=[{"role": "user", "content": prompt}],
    )

    raw = message.content[0].text.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    return json.loads(raw.strip())


# ── Sidebar ───────────────────────────────────────────────────────────────────
icp = load_icp()
pipeline = load_pipeline()

with st.sidebar:
    _logo = Path(__file__).parent / "assets" / "wbr_logo.png"
    if _logo.exists():
        st.image(str(_logo), use_container_width=True)
    else:
        st.markdown(
            '<div class="wbr-logo">WBR<span class="dot">.</span></div>'
            '<div class="wbr-sub">Worldwide Business Research</div>',
            unsafe_allow_html=True,
        )
    st.markdown(f"### {_event_cfg.get('short_name','WBR')} SDR Agent")
    st.caption(f"{_event_cfg.get('name','')} · {_event_cfg.get('location','')} · {_event_cfg.get('dates','')}")
    st.divider()

    if icp:
        st.metric("Registered Buyers", icp["buyer_count"])
        st.metric("Senior (VP+/Director)", f"{icp['senior_buyer_pct']}%")
        st.metric("Companies in Pipeline", len(pipeline))
        approved = len([c for c in pipeline if c.get("status") == "approved"])
        st.metric("Emails Approved", approved)
    else:
        st.warning("Run icp_profile.py first")

    st.divider()
    # Radar badge
    radar_finds = load_json("radar_finds.json", [])
    unreviewed = [c for c in radar_finds if not c.get("reviewed")]
    if unreviewed:
        st.error(f"Radar: {len(unreviewed)} new finds waiting!")
    else:
        st.caption("Radar: no new finds")

    st.divider()
    # Storage status
    if storage.github_configured():
        st.success("Saving to GitHub (persistent)")
        if st.button("Refresh data"):
            storage.refresh_cache(event_id=_event_id)
            st.rerun()
    elif storage.gsheets_configured():
        st.success("Saving to Google Sheets")
        if st.button("Refresh data"):
            storage.refresh_cache(event_id=_event_id)
            st.rerun()
    else:
        st.warning("Local storage only\n(edits reset on reboot)")

    # Rebuild ICP from an uploaded attendee CSV (persists to storage)
    with st.expander("Rebuild ICP"):
        st.caption("Upload your attendee registration CSV to regenerate the buyer profile.")
        icp_csv = st.file_uploader("Attendee CSV", type="csv", key="icp_csv")
        if icp_csv and st.button("Rebuild ICP from CSV"):
            import tempfile, os
            from icp_profile import build_icp
            with tempfile.NamedTemporaryFile(delete=False, suffix=".csv", mode="wb") as tmp:
                tmp.write(icp_csv.read())
                tmp_path = tmp.name
            try:
                new_icp = build_icp(tmp_path)
                storage.save_icp(new_icp, event_id=_event_id)
                st.success(f"ICP rebuilt -- {new_icp['buyer_count']} buyers.")
                st.rerun()
            except Exception as e:
                st.error(f"Failed: {e}")
            finally:
                os.unlink(tmp_path)

    st.divider()
    # Outlook status
    if outlook.is_configured():
        if outlook.is_authenticated():
            profile = outlook.get_profile()
            name = profile.get("displayName", "") if profile else ""
            st.success(f"Outlook: {name or 'Connected'}")
        else:
            st.warning("Outlook: needs auth")
    else:
        st.caption("Outlook: awaiting Azure setup")

    st.divider()
    # Hot leads badge
    hot = [c for c in pipeline if c.get("priority") == "hot" and c.get("status") not in ("closed_won","closed_lost")]
    replied = [c for c in pipeline if c.get("status") == "replied"]
    if replied:
        st.error(f"{len(replied)} replied -- follow up!")
    if hot:
        st.warning(f"{len(hot)} hot prospects")

    st.divider()
    st.caption("Sales stages")
    st.markdown("Researched > Contacted > Replied > Meeting > Contract > Won / Lost")

# ── Helpers: activity log & CRM ───────────────────────────────────────────────
from datetime import date, datetime

STATUSES = [
    "researched", "contacted", "replied",
    "meeting_booked", "contract_out", "closed_won", "closed_lost",
]
STATUS_LABELS = {
    "researched":    "Researched",
    "contacted":     "Contacted",
    "replied":       "Replied",
    "meeting_booked":"Meeting Booked",
    "contract_out":  "Contract Out",
    "closed_won":    "Closed Won",
    "closed_lost":   "Closed Lost",
    # legacy values
    "contact_found": "Contacted",
    "approved":      "Contacted",
    "sent":          "Contacted",
    "skipped":       "Researched",
}
PRIORITY_LABELS = {"hot": "Hot", "medium": "Medium", "cold": "Cold"}

def log_activity(company: dict, activity_type: str, **kwargs) -> dict:
    """Append an activity entry to company['activity_log']."""
    if "activity_log" not in company:
        company["activity_log"] = []
    entry = {
        "type": activity_type,
        "date": date.today().isoformat(),
        "source": "manual",
        **kwargs,
    }
    company["activity_log"].append(entry)
    return company

def days_since_last_activity(company: dict) -> int | None:
    log = company.get("activity_log", [])
    if not log:
        return None
    last = max(log, key=lambda x: x.get("date",""))
    try:
        delta = date.today() - date.fromisoformat(last["date"])
        return delta.days
    except Exception:
        return None

def render_activity_log(company: dict, key_prefix: str):
    """Render timeline + manual log controls for a company."""
    log = sorted(company.get("activity_log", []), key=lambda x: x.get("date",""), reverse=True)

    icons = {
        "email_sent":       "Email Sent",
        "reply_received":   "Reply",
        "call":             "Call",
        "note":             "Note",
        "meeting_booked":   "Meeting",
        "contract_sent":    "Contract",
        "status_change":    "Stage Change",
    }

    if log:
        st.markdown("**Activity Log**")
        for entry in log:
            icon = icons.get(entry["type"], entry["type"])
            date_str = entry.get("date","")
            src = " (Outlook)" if entry.get("source") == "outlook" else ""
            detail = entry.get("subject") or entry.get("note") or entry.get("preview","")
            st.markdown(f"**{icon}**{src} - {date_str}")
            if detail:
                st.caption(f"  {detail[:120]}")
    else:
        st.caption("No activity logged yet.")

    st.markdown("**Log Activity**")
    log_cols = st.columns(4)
    if log_cols[0].button("Email Sent", key=f"{key_prefix}_log_sent"):
        touch_opt = st.session_state.get(f"{key_prefix}_touch", 1)
        company = log_activity(company, "email_sent", touch=touch_opt, subject=f"Touch {touch_opt}")
        pipeline_local = load_pipeline()
        pipeline_local = upsert_company(pipeline_local, company)
        if company.get("status") in ("researched",):
            company["status"] = "contacted"
        save_pipeline(pipeline_local)
        st.rerun()
    if log_cols[1].button("Reply In", key=f"{key_prefix}_log_reply"):
        company = log_activity(company, "reply_received")
        company["status"] = "replied"
        pipeline_local = load_pipeline()
        pipeline_local = upsert_company(pipeline_local, company)
        save_pipeline(pipeline_local)
        st.rerun()
    if log_cols[2].button("Mtg Booked", key=f"{key_prefix}_log_mtg"):
        company = log_activity(company, "meeting_booked")
        company["status"] = "meeting_booked"
        pipeline_local = load_pipeline()
        pipeline_local = upsert_company(pipeline_local, company)
        save_pipeline(pipeline_local)
        st.rerun()
    if log_cols[3].button("Call", key=f"{key_prefix}_log_call"):
        company = log_activity(company, "call")
        pipeline_local = load_pipeline()
        pipeline_local = upsert_company(pipeline_local, company)
        save_pipeline(pipeline_local)
        st.rerun()

    note_text = st.text_input("Add a note", key=f"{key_prefix}_note_input", placeholder="e.g. Left voicemail, asked about budget...")
    if st.button("Save Note", key=f"{key_prefix}_save_note") and note_text:
        company = log_activity(company, "note", note=note_text)
        pipeline_local = load_pipeline()
        pipeline_local = upsert_company(pipeline_local, company)
        save_pipeline(pipeline_local)
        st.rerun()

# ── Account page helpers ──────────────────────────────────────────────────────
def open_account(company_name: str):
    """Navigate to a company's account detail page."""
    st.session_state["view"] = "account"
    st.session_state["selected_company"] = company_name
    st.rerun()

def get_contacts(company: dict) -> list:
    """Return contacts list, migrating legacy single-contact fields in-memory."""
    contacts = company.get("contacts") or []
    if not contacts and (company.get("contact_name") or company.get("contact_email")):
        contacts = [{
            "name": company.get("contact_name", ""),
            "title": company.get("contact_title", ""),
            "email": company.get("contact_email", ""),
            "phone": "",
            "notes": "",
            "activity_log": [],
        }]
    return contacts

CONTACT_ACT_ICONS = {
    "email_sent": "Email", "reply_received": "Reply", "call": "Call",
    "meeting": "Meeting", "note": "Note", "task": "Task",
}

def log_contact_activity(contact: dict, atype: str, **kw):
    contact.setdefault("activity_log", [])
    contact["activity_log"].append({"type": atype, "date": date.today().isoformat(), "source": "manual", **kw})

def render_account_page(company_name: str):
    pipeline = load_pipeline()
    company = get_company(pipeline, company_name)

    if st.button("<- Back to pipeline"):
        st.session_state["view"] = None
        st.session_state.pop("selected_company", None)
        st.rerun()

    if not company:
        st.error(f"Account '{company_name}' not found.")
        return

    score = company.get("score", 0)
    score_color = "High" if score >= 80 else "Medium" if score >= 60 else "Low"
    st.title(f"{company['company']} (Score: {score}/100)")

    m = st.columns(4)
    m[0].metric("Score", f"{score}/100")
    m[1].metric("Tier", company.get("tier", "?"))
    m[2].metric("Priority", PRIORITY_LABELS.get(company.get("priority", ""), "—"))
    m[3].metric("Stage", STATUS_LABELS.get(company.get("status", ""), "—"))

    # Editable priority + stage
    e1, e2, e3 = st.columns([2, 2, 1])
    cur_pri = company.get("priority", "cold")
    cur_pri = cur_pri if cur_pri in ("hot", "medium", "cold") else "cold"
    new_pri = e1.selectbox("Priority", ["hot", "medium", "cold"], index=["hot", "medium", "cold"].index(cur_pri), key="acc_pri")
    cur_status = company.get("status", "researched")
    new_status = e2.selectbox("Sales stage", STATUSES, index=STATUSES.index(cur_status) if cur_status in STATUSES else 0, key="acc_stage")
    e3.write("")
    e3.write("")
    if e3.button("Save"):
        old = company.get("status", "")
        company["priority"] = new_pri
        company["status"] = new_status
        if old != new_status:
            company = log_activity(company, "status_change", note=f"{old} -> {new_status}")
        pipeline = upsert_company(pipeline, company)
        save_pipeline(pipeline)
        st.success("Saved")
        st.rerun()

    if company.get("what_they_do"):
        st.write(f"**What they do:** {company['what_they_do']}")
    if company.get("pitch_angle") or company.get("outreach_note"):
        st.info(f"**Pitch angle:** {company.get('pitch_angle', company.get('outreach_note', ''))}")
    if company.get("vs_sponsor"):
        st.warning(f"**Competitors already sponsoring:** {company['vs_sponsor']}")

    # -- Contacts --
    st.divider()
    st.subheader("Contacts")
    contacts = get_contacts(company)

    # Auto-save any contacts that exist but aren't yet persisted (e.g. pulled from Tiga)
    if contacts and not company.get("contacts"):
        company["contacts"] = contacts
        pipeline = upsert_company(pipeline, company)
        save_pipeline(pipeline)

    if not contacts:
        st.caption("No contacts yet. Add one below.")

    def _save_contact_field(ci, contacts, company, pipeline):
        """Auto-save callback -- fires when any contact field loses focus."""
        c = contacts[ci]
        c["name"]  = st.session_state.get(f"ct_name_{ci}",  c.get("name", ""))
        c["title"] = st.session_state.get(f"ct_title_{ci}", c.get("title", ""))
        c["email"] = st.session_state.get(f"ct_email_{ci}", c.get("email", ""))
        c["phone"] = st.session_state.get(f"ct_phone_{ci}", c.get("phone", ""))
        c["notes"] = st.session_state.get(f"ct_notes_{ci}", c.get("notes", ""))
        company["contacts"] = contacts
        updated = upsert_company(pipeline, company)
        save_pipeline(updated)

    for ci, contact in enumerate(contacts):
        label = contact.get("name") or "(unnamed contact)"
        if contact.get("title"):
            label += f" -- {contact['title']}"
        with st.expander(label, expanded=len(contacts) == 1):
            cc1, cc2 = st.columns(2)
            cc1.text_input("Name",  contact.get("name", ""),  key=f"ct_name_{ci}",
                           on_change=_save_contact_field, args=(ci, contacts, company, pipeline))
            cc2.text_input("Title", contact.get("title", ""), key=f"ct_title_{ci}",
                           on_change=_save_contact_field, args=(ci, contacts, company, pipeline))
            cc1.text_input("Email", contact.get("email", ""), key=f"ct_email_{ci}",
                           on_change=_save_contact_field, args=(ci, contacts, company, pipeline))
            cc2.text_input("Phone", contact.get("phone", ""), key=f"ct_phone_{ci}",
                           on_change=_save_contact_field, args=(ci, contacts, company, pipeline))
            st.text_area("Notes about this contact", contact.get("notes", ""), key=f"ct_notes_{ci}",
                         height=90, on_change=_save_contact_field, args=(ci, contacts, company, pipeline))

            if st.button("Remove", key=f"ct_del_{ci}"):
                contacts.pop(ci)
                company["contacts"] = contacts
                pipeline = upsert_company(pipeline, company)
                save_pipeline(pipeline)
                st.rerun()

            # Per-contact history
            st.markdown("**History**")
            clog = sorted(contact.get("activity_log", []), key=lambda x: x.get("date", ""), reverse=True)
            if clog:
                for entry in clog:
                    icon = CONTACT_ACT_ICONS.get(entry["type"], entry["type"])
                    detail = entry.get("note") or entry.get("subject") or entry.get("preview", "")
                    src = " (Outlook)" if entry.get("source") == "outlook" else ""
                    st.caption(f"{icon}{src} - {entry.get('date', '')} {('-- ' + detail) if detail else ''}")
            else:
                st.caption("No history yet.")

            # Log activity for this contact (manual; Azure will auto-add later)
            a = st.columns(4)
            for col, (atype, lbl) in zip(a, [("email_sent", "Email"), ("reply_received", "Reply"), ("call", "Call"), ("meeting", "Meeting")]):
                if col.button(lbl, key=f"ct_act_{ci}_{atype}"):
                    log_contact_activity(contact, atype)
                    company["contacts"] = contacts
                    pipeline = upsert_company(pipeline, company)
                    save_pipeline(pipeline)
                    st.rerun()
            tnote = st.text_input("Log a note/task for this contact", key=f"ct_actnote_{ci}", placeholder="e.g. Sent follow-up, scheduling call next week")
            if st.button("Add note", key=f"ct_addnote_{ci}") and tnote:
                log_contact_activity(contact, "note", note=tnote)
                company["contacts"] = contacts
                pipeline = upsert_company(pipeline, company)
                save_pipeline(pipeline)
                st.rerun()

    # Add a new contact
    with st.expander("Add a contact"):
        n1, n2 = st.columns(2)
        nn = n1.text_input("Name", key="newc_name")
        nt = n2.text_input("Title", key="newc_title")
        ne = n1.text_input("Email", key="newc_email")
        nph = n2.text_input("Phone", key="newc_phone")
        if st.button("Add contact", type="primary") and (nn or ne):
            contacts.append({"name": nn, "title": nt, "email": ne, "phone": nph, "notes": "", "activity_log": []})
            company["contacts"] = contacts
            pipeline = upsert_company(pipeline, company)
            save_pipeline(pipeline)
            st.success("Contact added")
            st.rerun()

    # -- Account-level tasks & activity --
    st.divider()
    st.subheader("Account Activity & Tasks")
    st.caption("Log touches and tasks manually now. Once Azure is connected, sent/received emails will auto-log here.")
    render_activity_log(company, key_prefix="acc")


# ── Routing: account detail view replaces the tabs when an account is open ──
if st.session_state.get("view") == "account" and st.session_state.get("selected_company"):
    render_account_page(st.session_state["selected_company"])
    st.stop()

# ── Hero banner ───────────────────────────────────────────────────────────────
st.markdown(
    f'<div class="hero">'
    f'<h1>{_event_cfg.get("name","WBR")} -- Sponsorship Pipeline</h1>'
    f'<p>Research, score, sequence, track -- for the {_event_cfg.get("focus","event")} leaders in the room.</p>'
    f'<span class="pill">{_event_cfg.get("location","")} - {_event_cfg.get("dates","")}</span>'
    f'</div>',
    unsafe_allow_html=True,
)

# ── Tabs ──────────────────────────────────────────────────────────────────────
tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9, tab10, tab11 = st.tabs([
    "Research", "Pipeline", "Outreach Queue",
    "Radar", "Prospect", "Import", "Funnel", "Outlook", "Contacts", "Setup", "Event Info",
])


# ════════════════════════════════════════════════════════════════════════════
# TAB 1 -- RESEARCH
# ════════════════════════════════════════════════════════════════════════════
with tab1:  # Research
    st.header("Research a Company")
    st.caption("Enter any company name -- the agent will research them online, score their fit, and explain why they should sponsor.")

    col1, col2 = st.columns([3, 1])
    with col1:
        company_input = st.text_input("Company name", placeholder="e.g. ServiceMax, Aquant, TeamViewer...")
    with col2:
        st.write("")
        st.write("")
        research_btn = st.button("Research", type="primary", use_container_width=True)

    if research_btn and company_input:
        if not icp:
            st.error("Run icp_profile.py first to generate icp_summary.json")
        else:
            with st.spinner(f"Researching {company_input}..."):
                try:
                    result = research_company(company_input, icp)

                    # Score color
                    score = result.get("score", 0)
                    tier = result.get("tier", "C")
                    score_color = "High" if score >= 80 else "Medium" if score >= 60 else "Low"

                    st.divider()
                    col_a, col_b, col_c = st.columns(3)
                    col_a.metric("Fit Score", f"{score}/100")
                    col_b.metric("Tier", f"Tier {tier} ({score_color})")
                    col_c.metric("Status", "Researched")

                    st.subheader(result["company"])
                    st.write(f"**What they do:** {result.get('what_they_do','')}")
                    st.write(f"**Who they sell to:** {result.get('who_they_sell_to','')}")

                    st.info(f"**Why they fit:** {result.get('fit_reason','')}")
                    st.success(f"**Your pitch angle:** {result.get('pitch_angle','')}")
                    st.warning(f"**Watch out for:** {result.get('risk','')}")

                    # Add to pipeline
                    st.divider()
                    if st.button("Add to Pipeline", type="primary"):
                        pipeline = upsert_company(pipeline, result)
                        save_pipeline(pipeline)
                        st.success(f"{result['company']} added to your pipeline!")
                        st.rerun()

                    # Store in session for re-use
                    st.session_state["last_research"] = result

                except Exception as e:
                    st.error(f"Error: {e}")

    elif "last_research" in st.session_state:
        r = st.session_state["last_research"]
        st.info(f"Last researched: **{r['company']}** -- Score {r['score']}/100")


# ════════════════════════════════════════════════════════════════════════════
# TAB 2 -- PIPELINE
# ════════════════════════════════════════════════════════════════════════════
with tab2:
    st.header("My Sponsor Pipeline")

    pipeline = load_pipeline()

    if not pipeline:
        st.info("No companies yet -- research one in the Research tab to get started.")
    else:
        # Filters
        col_f1, col_f2, col_f3, col_f4 = st.columns(4)
        status_filter = col_f1.selectbox("Sales stage", ["All"] + STATUSES)
        priority_filter = col_f2.selectbox("Priority", ["All", "hot", "medium", "cold"])
        tier_filter = col_f3.selectbox("AI tier", ["All", "A", "B", "C"])
        categories = sorted(set(c.get("category","") for c in pipeline if c.get("category")))
        cat_filter = col_f4.selectbox("Category", ["All"] + categories)

        filtered = pipeline
        if status_filter != "All":
            filtered = [c for c in filtered if c.get("status") == status_filter]
        if priority_filter != "All":
            filtered = [c for c in filtered if c.get("priority") == priority_filter]
        if tier_filter != "All":
            filtered = [c for c in filtered if str(c.get("tier","")) == tier_filter]
        if cat_filter != "All":
            filtered = [c for c in filtered if c.get("category","") == cat_filter]

        # Search by company name (fast way to find one among hundreds)
        search = st.text_input("Search company", key="pl_search", placeholder="Type a company name...")
        if search:
            filtered = [c for c in filtered if search.lower() in c.get("company", "").lower()]

        # Sort: hot first, then by score
        priority_order = {"hot": 0, "medium": 1, "cold": 2, None: 3, "": 3}
        filtered_sorted = sorted(filtered, key=lambda x: (priority_order.get(x.get("priority"), 3), -x.get("score", 0)))

        # Cap how many cards render at once -- rendering 400+ expanders is very slow.
        PAGE = 25
        total_filtered = len(filtered_sorted)
        show_n = st.session_state.get("pl_show_n", PAGE)
        st.caption(f"Showing {min(show_n, total_filtered)} of {total_filtered} (filtered from {len(pipeline)} total). Use search/filters to narrow.")
        page = filtered_sorted[:show_n]

        for idx, company in enumerate(page):
            score = company.get("score", 0)
            tier = company.get("tier", "?")
            status = company.get("status", "researched")
            priority = company.get("priority", "")
            score_label = "High" if score >= 80 else "Medium" if score >= 60 else "Low"
            status_label = STATUS_LABELS.get(status, status.replace("_"," ").title())
            priority_label = PRIORITY_LABELS.get(priority, "")
            days = days_since_last_activity(company)
            days_str = f"· {days}d ago" if days is not None else ""

            header = f"**{company['company']}** -- {score}/100 ({score_label}) · {status_label}"
            if priority_label:
                header += f" · {priority_label}"
            if days_str:
                header += f" {days_str}"

            with st.expander(header):
                contacts = get_contacts(company)
                if contacts:
                    for ct in contacts[:2]:
                        name = ct.get("name","").strip()
                        title = ct.get("title","").strip()
                        email = ct.get("email","").strip()
                        parts = " · ".join(filter(None, [name, title, email]))
                        if parts:
                            st.caption(f"Contact: {parts}")
                else:
                    st.caption("No contact yet")
                if st.button("Open account page", key=f"open_acct_{idx}"):
                    open_account(company["company"])
                col1, col2 = st.columns(2)
                with col1:
                    st.write(f"**What they do:** {company.get('what_they_do','--')}")
                    st.write(f"**Who they sell to:** {company.get('who_they_sell_to', company.get('what_they_do','--'))}")
                    if company.get("fit_reason"):
                        st.write(f"**Fit reason:** {company.get('fit_reason')}")
                    if company.get("hq"):
                        st.write(f"**HQ:** {company.get('hq')}")
                with col2:
                    if company.get("pitch_angle") or company.get("outreach_note"):
                        st.info(f"**Pitch angle:** {company.get('pitch_angle', company.get('outreach_note','--'))}")
                    if company.get("vs_sponsor"):
                        st.warning(f"**Competitors already sponsoring:** {company.get('vs_sponsor')}")
                    if company.get("risk"):
                        st.error(f"**Watch out for:** {company.get('risk')}")

                st.divider()
                # Priority + stage controls
                ctrl1, ctrl2, ctrl3 = st.columns(3)
                new_priority = ctrl1.selectbox(
                    "Priority", ["hot","medium","cold"],
                    index=["hot","medium","cold"].index(company.get("priority","cold")) if company.get("priority") in ["hot","medium","cold"] else 2,
                    key=f"pri_{idx}",
                )
                new_status = ctrl2.selectbox(
                    "Sales stage", STATUSES,
                    index=STATUSES.index(status) if status in STATUSES else 0,
                    key=f"stage_{idx}",
                )
                if ctrl3.button("Save Stage", key=f"savestage_{idx}"):
                    old_status = company.get("status","")
                    company["priority"] = new_priority
                    company["status"] = new_status
                    if old_status != new_status:
                        company = log_activity(company, "status_change", note=f"{old_status} -> {new_status}")
                    pipeline = upsert_company(pipeline, company)
                    save_pipeline(pipeline)
                    st.success("Saved!")
                    st.rerun()

                st.divider()
                # Contact info
                st.write("**Contact Details**")
                c1, c2, c3 = st.columns(3)
                contact_name = c1.text_input("Contact name", value=company.get("contact_name",""), key=f"cn_{idx}")
                contact_title = c2.text_input("Title", value=company.get("contact_title",""), key=f"ct_{idx}")
                contact_email = c3.text_input("Email", value=company.get("contact_email",""), key=f"ce_{idx}")

                col_a, col_b, col_c = st.columns(3)
                if col_a.button("Save Contact", key=f"save_{idx}"):
                    company["contact_name"] = contact_name
                    company["contact_title"] = contact_title
                    company["contact_email"] = contact_email
                    if company.get("status") == "researched":
                        company["status"] = "contacted"
                    pipeline = upsert_company(pipeline, company)
                    save_pipeline(pipeline)
                    st.success("Contact saved!")
                    st.rerun()

                if col_b.button("Draft Meeting Email", key=f"gen_{idx}"):
                    if not contact_name:
                        st.warning("Add a contact name first")
                    else:
                        with st.spinner("Writing personalized email..."):
                            try:
                                email = generate_meeting_email(company, contact_name, contact_title, icp)
                                st.session_state[f"meeting_email_{company['company']}"] = email
                                # Also store a full 3-touch sequence for Outreach Queue
                                emails = generate_emails(company, contact_name, contact_title, icp)
                                company["emails"] = emails
                                if company.get("status") == "researched":
                                    company["status"] = "contacted"
                                pipeline = upsert_company(pipeline, company)
                                save_pipeline(pipeline)
                            except Exception as e:
                                st.error(f"Error generating email: {e}")

                # Show generated meeting email inline (persists across reruns via session state)
                draft = st.session_state.get(f"meeting_email_{company['company']}")
                if draft:
                    st.markdown("**Meeting Email Draft**")
                    draft_subject = st.text_input("Subject", value=draft["subject"], key=f"draft_subj_{idx}")
                    draft_body = st.text_area("Body", value=draft["body"], height=200, key=f"draft_body_{idx}")
                    dcol1, dcol2 = st.columns(2)
                    if dcol1.button("Looks good -- keep it", key=f"draft_keep_{idx}"):
                        st.session_state.pop(f"meeting_email_{company['company']}", None)
                        st.success("Saved! Find the full 3-touch sequence in Outreach Queue.")
                    if dcol2.button("Regenerate", key=f"draft_regen_{idx}"):
                        st.session_state.pop(f"meeting_email_{company['company']}", None)
                        st.rerun()

                if col_c.button("Remove", key=f"del_{idx}"):
                    pipeline = [c for c in pipeline if c["company"] != company["company"]]
                    save_pipeline(pipeline)
                    st.rerun()

                st.divider()
                render_activity_log(company, key_prefix=f"pl_{idx}")

        if show_n < total_filtered:
            if st.button(f"Show {min(PAGE, total_filtered - show_n)} more"):
                st.session_state["pl_show_n"] = show_n + PAGE
                st.rerun()

        st.divider()
        rows = []
        for c in pipeline:
            contacts = get_contacts(c)
            base = {
                "Company": c.get("company",""),
                "Score": c.get("score",""),
                "Tier": c.get("tier",""),
                "Status": c.get("status",""),
                "What They Do": c.get("what_they_do",""),
                "Pitch Angle": c.get("pitch_angle",""),
            }
            if contacts:
                for ct in contacts:
                    rows.append({**base,
                        "Contact Name": ct.get("name",""),
                        "Contact Title": ct.get("title",""),
                        "Contact Email": ct.get("email",""),
                        "Contact Phone": ct.get("phone",""),
                        "LinkedIn": ct.get("linkedin_url",""),
                    })
            else:
                rows.append({**base, "Contact Name":"", "Contact Title":"", "Contact Email":"", "Contact Phone":"", "LinkedIn":""})
        import io as _io
        _buf = _io.StringIO()
        _writer = csv.DictWriter(_buf, fieldnames=rows[0].keys())
        _writer.writeheader()
        _writer.writerows(rows)
        st.download_button("📥 Export Pipeline to CSV", data=_buf.getvalue(), file_name="pipeline_export.csv", mime="text/csv")


# ════════════════════════════════════════════════════════════════════════════
# TAB 3 -- OUTREACH QUEUE
# ════════════════════════════════════════════════════════════════════════════
with tab3:
    st.header("Outreach Queue")
    st.caption("Review and approve email sequences before anything gets sent.")

    pipeline = load_pipeline()
    companies_with_emails = [c for c in pipeline if c.get("emails")]

    if not companies_with_emails:
        st.info("No emails drafted yet. Go to Pipeline, add a contact, and click 'Generate Emails'.")
    else:
        for company in companies_with_emails:
            status = company.get("status", "")
            status_icon = "Approved" if status == "approved" else "Pending"

            with st.expander(f"**{company['company']}** · {company.get('contact_name','No contact')} · {status_icon}"):
                for email in company.get("emails", []):
                    st.markdown(f"#### Touch {email['touch']} -- {email['send_day']}")
                    st.markdown(f"**Subject:** {email['subject']}")
                    st.text_area(
                        "Body",
                        value=email["body"],
                        height=180,
                        key=f"email_{company['company']}_{email['touch']}",
                    )
                    st.divider()

                col1, col2, col3 = st.columns(3)
                if col1.button("Approve Sequence", key=f"approve_{company['company']}", type="primary"):
                    company["status"] = "approved"
                    pipeline = upsert_company(pipeline, company)
                    save_pipeline(pipeline)
                    st.success(f"{company['company']} approved! Ready to load into Tiga.")
                    st.rerun()

                if col2.button("Regenerate", key=f"regen_{company['company']}"):
                    with st.spinner("Rewriting..."):
                        emails = generate_emails(
                            company,
                            company.get("contact_name","[First Name]"),
                            company.get("contact_title",""),
                            icp
                        )
                        company["emails"] = emails
                        pipeline = upsert_company(pipeline, company)
                        save_pipeline(pipeline)
                        st.rerun()

                if col3.button("Skip", key=f"skip_{company['company']}"):
                    company["status"] = "skipped"
                    pipeline = upsert_company(pipeline, company)
                    save_pipeline(pipeline)
                    st.rerun()

        # Export approved to Tiga-ready CSV
        approved = [c for c in pipeline if c.get("status") == "approved"]
        if approved:
            st.divider()
            st.subheader(f"{len(approved)} sequences approved")
            if st.button("Export Approved to Tiga CSV", type="primary"):
                rows = []
                for c in approved:
                    emails = c.get("emails", [{},{},{}])
                    rows.append({
                        "First Name": c.get("contact_name","").split(" ")[0],
                        "Last Name": " ".join(c.get("contact_name","").split(" ")[1:]),
                        "Email": c.get("contact_email",""),
                        "Job Title": c.get("contact_title",""),
                        "Company": c.get("company",""),
                        "Email 1 Subject": emails[0].get("subject","") if len(emails) > 0 else "",
                        "Email 1 Body": emails[0].get("body","") if len(emails) > 0 else "",
                        "Email 2 Subject": emails[1].get("subject","") if len(emails) > 1 else "",
                        "Email 2 Body": emails[1].get("body","") if len(emails) > 1 else "",
                        "Email 3 Subject": emails[2].get("subject","") if len(emails) > 2 else "",
                        "Email 3 Body": emails[2].get("body","") if len(emails) > 2 else "",
                    })
                with open("tiga_import.csv", "w", newline="", encoding="utf-8") as f:
                    writer = csv.DictWriter(f, fieldnames=rows[0].keys())
                    writer.writeheader()
                    writer.writerows(rows)
                st.success("Saved to tiga_import.csv -- ready to import into Tiga!")


# ════════════════════════════════════════════════════════════════════════════
# TAB 4 -- RADAR
# ════════════════════════════════════════════════════════════════════════════
with tab4:
    st.header("Prospecting Radar")
    st.caption("New companies found automatically based on event signals, press releases, and competitor tracking. Review and approve before they hit your pipeline.")

    RADAR_FILE = "radar_finds.json"
    radar_finds = load_json(RADAR_FILE, [])
    unreviewed = [c for c in radar_finds if not c.get("reviewed")]
    reviewed = [c for c in radar_finds if c.get("reviewed")]

    # Manual run button
    col_r1, col_r2, col_r3 = st.columns([2, 1, 1])
    col_r1.info(f"**{len(unreviewed)} new companies** waiting for review · {len(reviewed)} previously reviewed")
    auto_add_toggle = col_r3.checkbox("Auto-add to pipeline", value=True, help="Automatically add finds (score >= 60) straight to your pipeline without manual review")
    if col_r2.button("Run Radar Now", type="primary"):
        with st.spinner("Searching queries across the web... this takes 10-15 minutes, keep this tab open"):
            try:
                import radar as radar_module
                importlib.reload(radar_module)
                finds = radar_module.run_radar(auto_add=auto_add_toggle, auto_add_min_score=60)
                added = len([f for f in finds if f.get("auto_added")])
                queued = len([f for f in finds if not f.get("auto_added")])
                msg = f"Found {len(finds)} new companies!"
                if auto_add_toggle and added:
                    msg += f" {added} added to pipeline, {queued} queued for review."
                st.success(msg)
                st.rerun()
            except Exception as e:
                st.error(f"Radar error: {e}")

    st.divider()

    if not unreviewed:
        st.success("All caught up -- no new companies to review.")
        st.caption("The radar runs every morning automatically. Check back tomorrow for new finds.")
    else:
        st.subheader(f"Review New Finds ({len(unreviewed)})")

        if st.button(f"Add all {len(unreviewed)} to Pipeline", type="primary", key="radar_add_all"):
            pipeline = load_pipeline()
            for rc in radar_finds:
                if not rc.get("reviewed"):
                    rc["reviewed"] = True
                    entry = {k: v for k, v in rc.items() if k not in ("reviewed", "dismissed")}
                    entry["status"] = "researched"
                    pipeline = upsert_company(pipeline, entry)
            save_pipeline(pipeline)
            with open(RADAR_FILE, "w") as f:
                json.dump(radar_finds, f, indent=2)
            st.success(f"Added {len(unreviewed)} companies to your pipeline!")
            st.rerun()

        for ridx, company in enumerate(unreviewed):
            score = company.get("score", 0)
            tier = company.get("tier", "?")
            score_label = "High" if score >= 80 else "Medium" if score >= 60 else "Low"
            signal = company.get("signal", "")

            with st.expander(f"**{company['company']}** -- {score}/100 · Tier {tier} · {signal}"):
                col1, col2 = st.columns(2)
                with col1:
                    st.write(f"**What they do:** {company.get('what_they_do','--')}")
                    st.write(f"**Category:** {company.get('category','--')}")
                    st.write(f"**Found:** {company.get('found_date','--')}")
                    if company.get("source_url"):
                        st.write(f"**Source:** [{company['source_url'][:60]}...]({company['source_url']})")
                with col2:
                    st.info(f"**Pitch angle:** {company.get('pitch_angle','--')}")
                    st.write(f"**Fit reason:** {company.get('fit_reason','--')}")

                col_a, col_b = st.columns(2)
                if col_a.button("Add to Pipeline", key=f"radar_add_{ridx}", type="primary"):
                    company["reviewed"] = True
                    company["status"] = "researched"
                    pipeline = upsert_company(pipeline, company)
                    save_pipeline(pipeline)
                    # Update radar file
                    for rc in radar_finds:
                        if rc["company"] == company["company"]:
                            rc["reviewed"] = True
                    with open(RADAR_FILE, "w") as f:
                        json.dump(radar_finds, f, indent=2)
                    st.success(f"{company['company']} added to pipeline!")
                    st.rerun()

                if col_b.button("Dismiss", key=f"radar_skip_{ridx}"):
                    for rc in radar_finds:
                        if rc["company"] == company["company"]:
                            rc["reviewed"] = True
                            rc["dismissed"] = True
                    with open(RADAR_FILE, "w") as f:
                        json.dump(radar_finds, f, indent=2)
                    st.rerun()

    # Previously reviewed
    if reviewed:
        with st.expander(f"Previously reviewed ({len(reviewed)})"):
            for c in reviewed:
                dismissed = c.get("dismissed", False)
                icon = "Dismissed" if dismissed else "Added"
                st.write(f"{icon}: {c['company']} -- {c.get('signal','')}")


# ════════════════════════════════════════════════════════════════════════════
# TAB 5 -- PROSPECT (Tiga Apollo + Signals)
# ════════════════════════════════════════════════════════════════════════════
with tab5:
    st.header("Prospect")
    st.caption("Search Seamless.ai using your ICP keywords, run your Tiga signals to score results, and auto-build a pipeline list.")

    import tiga_prospector as tp

    # -- Load signals --
    @st.cache_data(ttl=300, show_spinner=False)
    def fetch_signals():
        try:
            return tp.list_signals()
        except Exception as e:
            return []

    with st.spinner("Loading your Tiga signals..."):
        signals = fetch_signals()

    col_p1, col_p2 = st.columns([2, 1])

    with col_p1:
        st.subheader("Search settings")

        # Signal picker
        if signals:
            signal_options = {f"{tp._signal_name(s)} (ID: {s.get('id')})": s.get("id") for s in signals}
            selected_signal_labels = st.multiselect(
                "Score results with your Tiga signals (optional)",
                options=list(signal_options.keys()),
                help="Selected signals run against every company found. Leave blank to skip scoring."
            )
            selected_signal_ids = [str(signal_options[lbl]) for lbl in selected_signal_labels]
        else:
            st.info("No Tiga signals found -- results won't be scored. You can still build the list.")
            selected_signal_ids = []

        min_score = st.slider(
            "Only show results with signal score >=",
            min_value=0, max_value=100, value=0, step=5,
            help="0 = show everything. Requires at least one signal selected.",
            disabled=len(selected_signal_ids) == 0,
        )

        auto_add = st.checkbox(
            "Auto-add companies scoring >= 70 to pipeline",
            value=False,
            help="Adds high-confidence finds straight to pipeline without manual review"
        )

    with col_p2:
        st.subheader("Search preview")
        st.caption("Seamless searches these keyword combos:")
        st.markdown("""
- field service management software
- workforce management software
- predictive maintenance software
- IoT field service platform
- AR remote assistance field service
- service parts planning software
- knowledge management field technician
- AI service operations software
        """)
        st.caption("Edit tiga_prospector.py SEAMLESS_SEARCHES to customize.")

    st.divider()

    if st.button("Build Prospect List", type="primary", key="prospect_run"):
        prog_bar = st.progress(0)
        prog_label = st.empty()
        results_placeholder = st.empty()

        def _progress(label, current, total):
            prog_label.write(label)
            if total > 0:
                prog_bar.progress(min(current / total, 1.0))

        try:
            prospects = tp.run_prospecting(
                signal_ids=selected_signal_ids or None,
                min_score=min_score if selected_signal_ids else 0,
                auto_add=auto_add,
                auto_add_min_score=70,
                progress_cb=_progress,
            )
            prog_bar.empty()
            prog_label.empty()

            if not prospects:
                st.warning("No new companies found -- they may all already be in your pipeline.")
            else:
                auto_added = len([p for p in prospects if p.get("auto_added")])
                pending = len(prospects) - auto_added
                st.success(f"Found **{len(prospects)} new companies** · {auto_added} auto-added to pipeline · {pending} ready to review")

                # Store in session for review
                st.session_state["prospect_results"] = prospects
                st.rerun()

        except Exception as e:
            prog_bar.empty()
            prog_label.empty()
            st.error(f"Prospecting error: {e}")

    # -- Review results --
    prospects = st.session_state.get("prospect_results", [])
    if prospects:
        st.subheader(f"Results ({len(prospects)} companies)")

        not_added = [p for p in prospects if not p.get("auto_added")]
        added = [p for p in prospects if p.get("auto_added")]

        if added:
            st.success(f"{len(added)} already auto-added to pipeline")

        if not_added:
            if st.button(f"Add all {len(not_added)} to pipeline", type="primary", key="prospect_add_all"):
                pipeline = load_pipeline()
                for p in not_added:
                    entry = {
                        "company": p["company"],
                        "category": p.get("industry", "Other"),
                        "what_they_do": p.get("description", ""),
                        "score": p.get("signal_score", 0),
                        "tier": "B",
                        "priority": "medium",
                        "status": "researched",
                        "source": "tiga_prospector",
                        "website": p.get("website", ""),
                        "linkedin_url": p.get("linkedin_url", ""),
                        "fit_reason": p.get("signal_reasoning", "Found via Tiga Apollo prospecting"),
                        "pitch_angle": "",
                        "outreach_note": "",
                        "contacts": [],
                    }
                    pipeline = upsert_company(pipeline, entry)
                save_pipeline(pipeline)
                st.success(f"Added {len(not_added)} companies!")
                st.session_state["prospect_results"] = []
                st.rerun()

            for pidx, p in enumerate(not_added):
                sig_str = f" · Signal: {p['signal_score']}/100" if "signal_score" in p else ""
                score_label = "High" if p.get("signal_score", 0) >= 70 else "Medium" if p.get("signal_score", 0) >= 40 else "Low"
                with st.expander(f"**{p['company']}**{sig_str} · {p.get('industry','?')} · {p.get('headcount','?')} employees"):
                    col1, col2 = st.columns(2)
                    with col1:
                        st.write(f"**Website:** {p.get('website') or '--'}")
                        st.write(f"**Location:** {p.get('location') or '--'}")
                        if p.get("linkedin_url"):
                            st.write(f"**LinkedIn:** [{p['linkedin_url'][:50]}]({p['linkedin_url']})")
                    with col2:
                        if p.get("description"):
                            st.write(f"**About:** {p['description'][:200]}")
                        if p.get("signal_reasoning"):
                            st.info(f"**Signal:** {p['signal_reasoning'][:200]}")

                    if st.button("Add to Pipeline", key=f"prospect_add_{pidx}", type="primary"):
                        pipeline = load_pipeline()
                        entry = {
                            "company": p["company"],
                            "category": p.get("industry", "Other"),
                            "what_they_do": p.get("description", ""),
                            "score": p.get("signal_score", 0),
                            "tier": "B",
                            "priority": "medium",
                            "status": "researched",
                            "source": "tiga_prospector",
                            "website": p.get("website", ""),
                            "fit_reason": p.get("signal_reasoning", "Found via Tiga Apollo"),
                            "pitch_angle": "",
                            "outreach_note": "",
                            "contacts": [],
                        }
                        pipeline = upsert_company(pipeline, entry)
                        save_pipeline(pipeline)
                        p["auto_added"] = True
                        st.success(f"{p['company']} added!")
                        st.rerun()

with tab6:
    st.header("Import")
    st.caption("Upload a CSV from LinkedIn Sales Navigator. The agent auto-detects accounts vs. contacts, deduplicates against your pipeline, and enriches contacts via Tiga.")

    import salesnav_import as sni
    import io
    import pandas as pd

    # -- How-to guide --
    with st.expander("How to export from LinkedIn Sales Navigator"):
        st.markdown("""
**Accounts list (adds companies to pipeline):**
1. In Sales Nav, Accounts, run your search
2. Select accounts, Export, CSV
3. Upload here -- agent reads: Account Name, Industry, Headcount, Website, HQ

**Leads/Contacts list (finds people at your accounts):**
1. In Sales Nav, Leads, run your search
2. Select leads, Export, CSV
3. Upload here -- agent reads: First/Last Name, Title, Company, Email, LinkedIn URL

**Auto-detect:** The agent figures out which type it is automatically.
        """)

    st.divider()

    uploaded_file = st.file_uploader(
        "Drop your Sales Navigator CSV here",
        type=["csv"],
        help="Accounts export or Leads export from LinkedIn Sales Navigator"
    )

    if uploaded_file:
        import tempfile
        raw_bytes = uploaded_file.read()
        with tempfile.NamedTemporaryFile(delete=False, suffix=".csv", mode="wb") as tmp:
            tmp.write(raw_bytes)
            tmp_path = tmp.name

        import csv as _csv
        decoded = raw_bytes.decode("utf-8-sig", errors="replace")
        reader = _csv.DictReader(io.StringIO(decoded))
        headers = list(reader.fieldnames or [])
        csv_type = sni.detect_csv_type(headers)

        df_preview = pd.read_csv(io.StringIO(decoded), encoding="utf-8-sig", on_bad_lines="skip")
        type_label = {"accounts": "Accounts list", "contacts": "Contacts list", "unknown": "Unknown"}.get(csv_type, "Unknown")
        st.success(f"Loaded **{len(df_preview)} rows** · {type_label} detected")
        st.dataframe(df_preview.head(5), use_container_width=True)

        if csv_type == "unknown":
            st.warning("Could not auto-detect CSV type. Please check column names match Sales Nav export format.")

        elif csv_type == "accounts":
            st.subheader("Import Accounts to Pipeline")
            entries, skipped, total = sni.parse_accounts_csv(tmp_path)
            col_a1, col_a2 = st.columns(2)
            col_a1.metric("New companies", len(entries))
            col_a2.metric("Already in pipeline (skipped)", len(skipped))

            if entries:
                with st.expander(f"Preview {len(entries)} new companies"):
                    for e in entries[:20]:
                        st.write(f"- **{e['company']}** -- {e.get('category','') or 'No industry'} · {e.get('headcount','') or '?'} employees")
                    if len(entries) > 20:
                        st.caption(f"...and {len(entries)-20} more")

                col_opt1, col_opt2 = st.columns(2)
                score_them = col_opt1.checkbox("Score each company with AI", value=False,
                    help="Uses Claude to score sponsor fit -- ~1-2 min per 10 companies. You can always score later from the Pipeline tab.")
                limit = col_opt2.number_input("Max to import", min_value=1,
                    max_value=len(entries), value=min(100, len(entries)))

                if st.button("Add to Pipeline", type="primary", key="sn_accounts_import"):
                    to_add = entries[:limit]
                    if score_them:
                        prog = st.progress(0)
                        stat = st.empty()
                        for i, entry in enumerate(to_add):
                            stat.write(f"Scoring {entry['company']}... ({i+1}/{len(to_add)})")
                            try:
                                scored = research_company(entry["company"], icp)
                                entry.update({k: v for k, v in scored.items() if k not in ("company",)})
                                entry["source"] = "salesnav_import"
                            except Exception:
                                pass
                            prog.progress((i+1) / len(to_add))
                        stat.empty()
                        prog.empty()

                    n = sni.add_accounts_to_pipeline(to_add)
                    st.success(f"Added {n} companies to your pipeline!")
                    st.rerun()
            else:
                st.info("All companies in this file are already in your pipeline.")

        elif csv_type == "contacts":
            st.subheader("Import Contacts")
            contacts, total = sni.parse_contacts_csv(tmp_path)
            missing_email = len([c for c in contacts if not c.get("email")])
            has_email = len(contacts) - missing_email

            col_c1, col_c2, col_c3 = st.columns(3)
            col_c1.metric("Contacts parsed", len(contacts))
            col_c2.metric("Have email", has_email)
            col_c3.metric("Need enrichment", missing_email)

            if contacts:
                with st.expander(f"Preview {len(contacts)} contacts"):
                    for c in contacts[:20]:
                        email_str = c.get("email") or "no email"
                        st.write(f"- **{c['name']}** · {c.get('title','')} @ {c.get('company','')} · {email_str}")
                    if len(contacts) > 20:
                        st.caption(f"...and {len(contacts)-20} more")

                enrich = st.checkbox(
                    f"Enrich {missing_email} contacts missing email via Tiga",
                    value=missing_email > 0,
                    disabled=missing_email == 0,
                    help="Runs Tiga waterfall enrichment to find work emails for contacts that Sales Nav didn't include"
                )

                if st.button("Save Contacts", type="primary", key="sn_contacts_import"):
                    if enrich and missing_email:
                        prog = st.progress(0)
                        stat = st.empty()
                        def _cb(i, total_n, name):
                            stat.write(f"Enriching {name}... ({i+1}/{total_n})")
                            prog.progress((i+1) / max(total_n, 1))
                        contacts, n_enriched = sni.enrich_contacts_batch(contacts, progress_cb=_cb)
                        stat.empty()
                        prog.empty()
                        st.info(f"Enriched {n_enriched} contacts via Tiga")

                    n_saved = sni.save_contacts_csv(contacts)
                    st.success(f"Saved {n_saved} contacts to contacts.csv!")
                    st.caption("Go to the Contacts tab to view them.")

        # Cleanup temp file
        try:
            import os as _os
            _os.unlink(tmp_path)
        except Exception:
            pass


# ════════════════════════════════════════════════════════════════════════════
# TAB 7 -- FUNNEL
# ════════════════════════════════════════════════════════════════════════════
with tab7:
    st.header("Sales Funnel")
    st.caption("Pipeline by sales stage and priority.")

    pipeline = load_pipeline()

    if not pipeline:
        st.info("No companies yet -- start in Research.")
    else:
        # -- Summary metrics --
        m_cols = st.columns(7)
        for i, s in enumerate(STATUSES):
            count = sum(1 for c in pipeline if c.get("status") == s)
            m_cols[i].metric(STATUS_LABELS[s], count)

        st.divider()

        # -- Priority breakdown --
        st.subheader("By Priority")
        p_cols = st.columns(3)
        for col, (pkey, plabel) in zip(p_cols, PRIORITY_LABELS.items()):
            companies = [c for c in pipeline if c.get("priority") == pkey and c.get("status") not in ("closed_won","closed_lost")]
            with col:
                st.markdown(f"### {plabel} ({len(companies)})")
                ranked = sorted(companies, key=lambda x: -x.get("score",0))
                for ri, c in enumerate(ranked[:12]):
                    score = c.get("score",0)
                    status_label = STATUS_LABELS.get(c.get("status",""), c.get("status",""))
                    days = days_since_last_activity(c)
                    days_str = f" · {days}d" if days is not None else ""
                    if st.button(f"{c['company']}", key=f"fn_pri_{pkey}_{ri}"):
                        open_account(c["company"])
                    st.caption(f"{status_label}{days_str}")
                if len(ranked) > 12:
                    st.caption(f"+ {len(ranked) - 12} more")

        st.divider()

        # -- Stage columns --
        st.subheader("By Sales Stage")
        active_statuses = [s for s in STATUSES if s not in ("closed_won","closed_lost")]
        stage_cols = st.columns(len(active_statuses))
        for col, s in zip(stage_cols, active_statuses):
            companies = [c for c in pipeline if c.get("status") == s or (s == "researched" and c.get("status") not in STATUSES)]
            with col:
                st.caption(f"{STATUS_LABELS[s]} · {len(companies)}")
                for si, c in enumerate(sorted(companies, key=lambda x: -x.get("score",0))[:8]):
                    badge = PRIORITY_LABELS.get(c.get("priority",""), "")
                    contacts = get_contacts(c)
                    score = c.get("score", 0)

                    contact_info = ""
                    for ct in contacts[:1]:
                        name = ct.get("name", "").strip()
                        if name:
                            contact_info = f" · {name}"

                    st.markdown(f"**{c['company']}** {badge}{contact_info}")
                    st.caption(f"{score}/100")
                    if st.button("Open", key=f"fn_stg_{s}_{si}", use_container_width=True):
                        open_account(c["company"])

        st.divider()

        # -- Closed --
        won = [c for c in pipeline if c.get("status") == "closed_won"]
        lost = [c for c in pipeline if c.get("status") == "closed_lost"]
        w_col, l_col = st.columns(2)
        with w_col:
            st.markdown(f"### Closed Won ({len(won)})")
            for c in won:
                st.markdown(f"**{c['company']}**")
        with l_col:
            st.markdown(f"### Closed Lost ({len(lost)})")
            for c in lost:
                st.markdown(f"**{c['company']}**")

        st.divider()

        # -- Activity feed --
        st.subheader("Recent Activity")
        all_activities = []
        for c in pipeline:
            for entry in c.get("activity_log", []):
                all_activities.append({**entry, "_company": c["company"]})
        all_activities.sort(key=lambda x: x.get("date",""), reverse=True)

        if all_activities:
            for entry in all_activities[:20]:
                detail = entry.get("subject") or entry.get("note") or entry.get("preview","")
                st.markdown(f"**{entry['_company']}** · {entry['type'].replace('_',' ').title()} · {entry.get('date','')}")
                if detail:
                    st.caption(f"  {detail[:100]}")
        else:
            st.info("No activity logged yet. Use the Pipeline tab to log sends, replies, and notes.")

# ════════════════════════════════════════════════════════════════════════════
# TAB 8 -- OUTLOOK
# ════════════════════════════════════════════════════════════════════════════
with tab8:
    st.header("Outlook Integration")

    if not outlook.is_configured():
        st.info(
            "**Waiting on Azure App Registration.**\n\n"
            "Once Gabe provides the credentials, add them to your .env file on Streamlit Cloud "
            "(Settings > Secrets) and this tab activates automatically.\n\n"
            "AZURE_CLIENT_ID=...\nAZURE_CLIENT_SECRET=...\nAZURE_TENANT_ID=..."
        )
        st.divider()
        st.subheader("What this will do once connected")
        st.markdown(
            "- **Reply detection** -- flags when a prospect replies to your outreach\n"
            "- **Sent log** -- shows which touches have been sent per company\n"
            "- **Send from here** -- send approved sequences directly without going to Outlook\n"
            "- **Hot leads** -- surfaces replies at the top so nothing slips through"
        )

    elif not outlook.is_authenticated():
        st.warning("Outlook is configured but needs one-time authorization.")
        auth_url = outlook.get_auth_url()
        st.markdown(f"[Click here to authorize Outlook access]({auth_url})")
        st.caption("You'll be redirected back to the app. Paste the code= value from the URL below if it doesn't auto-complete.")
        code = st.text_input("Authorization code (from redirect URL)")
        if code and st.button("Complete Authorization", type="primary"):
            try:
                outlook.exchange_code_for_token(code)
                st.success("Connected!")
                st.rerun()
            except Exception as e:
                st.error(f"Authorization failed: {e}")

    else:
        profile = outlook.get_profile()
        st.success(f"Connected as **{profile.get('displayName', '')}** ({profile.get('mail', '')})")
        st.divider()

        pipeline = load_pipeline()
        prospect_emails = [c.get("contact_email", "") for c in pipeline if c.get("contact_email")]

        col1, col2 = st.columns(2)

        with col1:
            st.subheader("Replies from Prospects")
            with st.spinner("Checking inbox..."):
                replies = outlook.get_recent_replies(prospect_emails)
            if replies:
                for msg in replies:
                    sender = msg["from"]["emailAddress"]["address"]
                    match = next((c for c in pipeline if c.get("contact_email","").lower() == sender.lower()), None)
                    company_name = match["company"] if match else sender
                    st.markdown(f"**{company_name}** · {msg['receivedDateTime'][:10]}")
                    st.caption(f"Re: {msg['subject']}")
                    st.caption(msg["bodyPreview"][:150])
                    if match and st.button("Log Reply + Update Stage", key=f"reply_{msg['id'][:8]}"):
                        match = log_activity(match, "reply_received",
                                            subject=msg["subject"],
                                            preview=msg["bodyPreview"][:150],
                                            source="outlook")
                        match["status"] = "replied"
                        pipeline = upsert_company(pipeline, match)
                        save_pipeline(pipeline)
                        st.success("Logged and stage updated to Replied!")
                        st.rerun()
                    st.divider()
            else:
                st.info("No replies from prospects in recent inbox.")

        with col2:
            st.subheader("Sent to Prospects")
            with st.spinner("Checking sent items..."):
                sent = outlook.get_sent_to_prospects(prospect_emails)
            if sent:
                for msg in sent:
                    to_addr = msg["toRecipients"][0]["emailAddress"]["address"] if msg.get("toRecipients") else ""
                    match = next((c for c in pipeline if c.get("contact_email","").lower() == to_addr.lower()), None)
                    company_name = match["company"] if match else to_addr
                    st.markdown(f"**{company_name}**")
                    st.caption(f"{msg['subject']} · {msg['sentDateTime'][:10]}")
                    st.divider()
            else:
                st.info("No sent emails to prospects found.")

        st.divider()
        st.subheader("Send an Email")
        contacts_with_email = [c for c in pipeline if c.get("contact_email") and c.get("emails")]
        if not contacts_with_email:
            st.info("No contacts with both an email address and generated sequences yet.")
        else:
            options = {f"{c['company']} -- {c.get('contact_name', c['contact_email'])}": c for c in contacts_with_email}
            selected_label = st.selectbox("Select contact", list(options.keys()))
            selected_company = options[selected_label]

            touch_num = st.selectbox("Which touch", ["Touch 1 (Day 1)", "Touch 2 (Day 4)", "Touch 3 (Day 9)"])
            touch_idx = int(touch_num[6]) - 1
            email_data = selected_company["emails"][touch_idx]

            subject = st.text_input("Subject", value=email_data["subject"])
            body = st.text_area("Body", value=email_data["body"], height=220)

            if st.button("Send", type="primary"):
                success = outlook.send_email(selected_company["contact_email"], subject, body)
                if success:
                    selected_company = log_activity(selected_company, "email_sent",
                                                    touch=touch_idx + 1,
                                                    subject=subject,
                                                    source="outlook")
                    if selected_company.get("status") == "researched":
                        selected_company["status"] = "contacted"
                    pipeline = upsert_company(pipeline, selected_company)
                    save_pipeline(pipeline)
                    st.success(f"Sent to {selected_company['contact_email']}")
                    st.rerun()
            else:
                st.error("Send failed -- check your Outlook connection.")


# ════════════════════════════════════════════════════════════════════════════
# TAB 9 -- CONTACTS
# ════════════════════════════════════════════════════════════════════════════
with tab9:
    try:
        import pandas as pd
        import traceback

        st.header("All Contacts")
        st.caption("Every contact across all pipeline accounts, in one place.")

        if st.button("Refresh"):
            storage.refresh_cache(event_id=_event_id)
            st.rerun()

        pipeline_data = load_pipeline()

        all_rows = []
        for company in pipeline_data:
            for contact in get_contacts(company):
                all_rows.append({
                    "Company": company.get("company", ""),
                    "Tier":    company.get("tier", ""),
                    "Score":   company.get("score", ""),
                    "Status":  company.get("status", ""),
                    "Name":    contact.get("name", ""),
                    "Title":   contact.get("title", ""),
                    "Email":   contact.get("email", ""),
                    "Phone":   contact.get("phone", ""),
                    "LinkedIn": contact.get("linkedin", ""),
                    "Notes":   contact.get("notes", ""),
                    "Source":  contact.get("source", ""),
                })

        if not all_rows:
            st.info("No contacts yet. Add contacts via the Pipeline tab or run tiga_contacts.py.")
        else:
            m1, m2, m3 = st.columns(3)
            m1.metric("Total Contacts", len(all_rows))
            m2.metric("Companies with Contacts", len({r["Company"] for r in all_rows}))
            m3.metric("Contacts with Email", sum(1 for r in all_rows if r["Email"]))

            st.divider()

            fc1, fc2, fc3 = st.columns(3)
            with fc1:
                filter_company = st.text_input("Filter by company", placeholder="Type to search...")
            with fc2:
                all_tiers = sorted({str(r["Tier"]) for r in all_rows if r["Tier"] != ""})
                filter_tier = st.selectbox("Tier", ["All"] + all_tiers)
            with fc3:
                filter_email_only = st.checkbox("Email only", value=False)

            filtered = all_rows
            if filter_company:
                filtered = [r for r in filtered if filter_company.lower() in r["Company"].lower()]
            if filter_tier != "All":
                filtered = [r for r in filtered if str(r["Tier"]) == filter_tier]
            if filter_email_only:
                filtered = [r for r in filtered if r["Email"]]

            st.caption(f"Showing {len(filtered)} of {len(all_rows)} contacts")

            display_cols = ["Company", "Tier", "Score", "Name", "Title", "Email", "Phone", "Status"]
            df = pd.DataFrame(filtered)[display_cols]
            st.dataframe(df, use_container_width=True, hide_index=True)

            st.divider()
            st.download_button(
                label="Download CSV",
                data=pd.DataFrame(filtered).to_csv(index=False),
                file_name="contacts_export.csv",
                mime="text/csv",
                type="primary",
            )

    except Exception as e:
        st.error(f"Contacts tab error: {e}")
        st.code(traceback.format_exc())


# ════════════════════════════════════════════════════════════════════════════
# TAB 10 -- SETUP  (only active for new/empty events)
# ════════════════════════════════════════════════════════════════════════════
with tab10:
    import event_setup
    import tempfile, os as _os

    pipeline_now = load_pipeline()
    icp_now      = load_icp()

    if pipeline_now:
        # Already set up -- show summary
        st.success(f"Setup complete. {len(pipeline_now)} companies in pipeline.")
        if icp_now:
            st.info(
                f"ICP: **{icp_now.get('buyer_count','?')} buyers** · "
                f"**{icp_now.get('senior_buyer_pct','?')}% senior** · "
                f"event focus: {_event_cfg.get('focus','')}"
            )
        st.caption("To reconfigure, use the Rebuild ICP option in the sidebar.")
    else:
        st.header("Event Setup Wizard")
        st.caption(
            f"Welcome to **{_event_cfg.get('name','')}**. "
            "Complete the two steps below to seed your pipeline."
        )

        # ── STEP 1: Attendee CSV ──────────────────────────────────────────
        st.divider()
        st.subheader("Step 1 — Upload Attendee Registration CSV")
        st.caption(
            "Export the registration list from your event platform. "
            "Required columns (same format as WBR events): "
            "**Account**, **Job Title**, **Price List Type** (Primary = buyer, Vendor = sponsor)."
        )

        attendee_csv = st.file_uploader(
            "Attendee CSV", type=["csv"], key="setup_attendee_csv"
        )
        icp_built = None
        if attendee_csv:
            raw_bytes = attendee_csv.read()
            if st.button("Build Buyer Profile from CSV", type="primary", key="setup_build_icp"):
                with st.spinner("Analysing attendees..."):
                    try:
                        with tempfile.NamedTemporaryFile(
                            delete=False, suffix=".csv", mode="wb"
                        ) as tmp:
                            tmp.write(raw_bytes)
                            tmp_path = tmp.name
                        icp_built = event_setup.build_icp_from_csv(tmp_path)
                        storage.save_icp(icp_built, event_id=_event_id)
                        _os.unlink(tmp_path)
                        st.success(
                            f"ICP built: **{icp_built['buyer_count']} buyers** · "
                            f"**{icp_built['senior_buyer_pct']}% VP/Director+**"
                        )
                        st.session_state["setup_icp_done"] = True
                        st.rerun()
                    except Exception as e:
                        st.error(f"CSV error: {e}")

        if st.session_state.get("setup_icp_done") or icp_now:
            st.success("Buyer profile ready.")

        # ── STEP 2: Event website ─────────────────────────────────────────
        st.divider()
        st.subheader("Step 2 — Research Event Website")
        st.caption(
            "Paste your event URL. The agent will scrape it and suggest "
            "the best search keywords for finding sponsor prospects."
        )

        default_url = _event_cfg.get("website", "")
        event_url = st.text_input(
            "Event website URL", value=default_url, key="setup_event_url"
        )

        scraped = st.session_state.get("setup_scraped")

        if st.button("Research Event", type="primary", key="setup_scrape_btn"):
            if not event_url:
                st.warning("Paste the event URL first.")
            else:
                with st.spinner("Researching event website..."):
                    try:
                        scraped = event_setup.research_event_website(
                            event_url, _event_cfg.get("name", "")
                        )
                        st.session_state["setup_scraped"] = scraped
                    except Exception as e:
                        st.error(f"Research error: {e}")

        if scraped:
            st.markdown("**What we found:**")
            st.info(f"**Event focus:** {scraped.get('focus','')}")

            col_t, col_s = st.columns(2)
            with col_t:
                st.markdown("**Key topics:**")
                for t in scraped.get("key_topics", []):
                    st.markdown(f"- {t}")
            with col_s:
                st.markdown("**Sponsor categories:**")
                for c in scraped.get("sponsor_categories", []):
                    st.markdown(f"- {c}")

            st.markdown("**Suggested search keywords:**")
            kw_val = st.text_area(
                "Edit if needed",
                value=scraped.get("suggested_search_keywords", ""),
                height=80,
                key="setup_keywords",
            )
            st.caption(
                "These keywords drive the Prospect and Radar tabs. "
                "They are saved to your event config — update events_registry.py "
                "with the final keywords to make them permanent."
            )

            if st.session_state.get("setup_icp_done") or icp_now:
                st.divider()
                st.subheader("Step 3 — Run First Radar Pass")
                st.caption(
                    "The agent will search for companies matching your event "
                    "audience and add the best fits to your pipeline."
                )
                if st.button(
                    "Seed Pipeline with Radar", type="primary", key="setup_radar_btn"
                ):
                    with st.spinner(
                        "Searching for sponsor prospects... this takes a few minutes."
                    ):
                        try:
                            import radar as radar_module
                            import importlib
                            importlib.reload(radar_module)
                            # Temporarily override search keywords from scraped data
                            kw_override = st.session_state.get("setup_keywords", "")
                            if kw_override:
                                _event_cfg["search_keywords"] = kw_override
                                st.session_state["event_cfg"] = _event_cfg
                            finds = radar_module.run_radar(
                                auto_add=True, auto_add_min_score=60
                            )
                            st.success(
                                f"Done! Found **{len(finds)} companies**. "
                                "Head to the Pipeline tab to review them."
                            )
                            st.session_state.pop("setup_scraped", None)
                            st.rerun()
                        except Exception as e:
                            st.error(f"Radar error: {e}")
            else:
                st.info("Complete Step 1 first to enable the radar seed.")


# ════════════════════════════════════════════════════════════════════════════
# TAB 11 -- EVENT INFO  (Agenda + Attendee List)
# ════════════════════════════════════════════════════════════════════════════
with tab11:
    import json as _json_ei
    import tempfile as _tmp_ei
    import os as _os_ei
    import io as _io_ei
    import csv as _csv_ei

    # ── helpers ──────────────────────────────────────────────────────────────
    _AGENDA_FILE = Path(__file__).parent / f"{_event_id}_agenda.json"

    def _load_agenda():
        if _AGENDA_FILE.exists():
            with open(_AGENDA_FILE) as _f:
                return _json_ei.load(_f)
        return []

    def _save_agenda(sessions):
        with open(_AGENDA_FILE, "w") as _f:
            _json_ei.dump(sessions, _f, indent=2)

    st.markdown(
        f'<div class="hero" style="padding:18px 24px;margin-bottom:10px;">'
        f'<h1 style="font-size:1.3rem;">{_event_cfg.get("name","Event")} — Event Info</h1>'
        f'<p>{_event_cfg.get("location","")} · {_event_cfg.get("dates","")}</p>'
        f'</div>',
        unsafe_allow_html=True,
    )

    ei_agenda_tab, ei_attendee_tab = st.tabs(["📋 Agenda", "👥 Attendee List"])

    # ── AGENDA ────────────────────────────────────────────────────────────────
    with ei_agenda_tab:
        st.subheader("Event Agenda")
        st.caption("Upload a CSV or add sessions manually. Stored per event so your whole team can see it.")

        sessions = _load_agenda()

        # ── Upload agenda CSV ──
        st.markdown("**Upload agenda file**")
        st.caption("Accepts CSV, Excel (.xlsx), or Word (.docx) — any layout works.")
        agenda_file = st.file_uploader(
            "Drag & drop your agenda here, or click to browse  (CSV, Excel, or Word)",
            key="agenda_upload",
        )
        if agenda_file:
            import pandas as _pd_ag
            fname = agenda_file.name.lower()
            df_ag = None
            parse_error = None

            try:
                if fname.endswith(".csv"):
                    raw = agenda_file.read().decode("utf-8-sig", errors="replace")
                    df_ag = _pd_ag.read_csv(_io_ei.StringIO(raw), on_bad_lines="skip")

                elif fname.endswith((".xlsx", ".xls")):
                    df_ag = _pd_ag.read_excel(agenda_file, engine="openpyxl")

                elif fname.endswith(".docx"):
                    from docx import Document as _DocxDoc
                    doc = _DocxDoc(agenda_file)
                    # Try the largest table first (most agendas live in a table)
                    best_table = max(doc.tables, key=lambda t: len(t.rows)) if doc.tables else None
                    if best_table and len(best_table.rows) > 1:
                        all_rows = []
                        for tr in best_table.rows:
                            all_rows.append([cell.text.strip() for cell in tr.cells])
                        # Normalize row lengths to match the widest row
                        max_cols = max(len(r) for r in all_rows)
                        all_rows = [r + [""] * (max_cols - len(r)) for r in all_rows]
                        headers = all_rows[0]
                        # Deduplicate blank/duplicate headers
                        seen = {}
                        clean_headers = []
                        for h in headers:
                            h = h or "Col"
                            seen[h] = seen.get(h, 0) + 1
                            clean_headers.append(h if seen[h] == 1 else f"{h}_{seen[h]}")
                        df_ag = _pd_ag.DataFrame(all_rows[1:], columns=clean_headers)
                    else:
                        # Fall back: each non-empty paragraph becomes a session row
                        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                        df_ag = _pd_ag.DataFrame({"Session": paras})

            except Exception as _pe:
                parse_error = str(_pe)

            if parse_error:
                st.error(f"Could not read file: {parse_error}")
            elif df_ag is not None and not df_ag.empty:
                df_ag.columns = [str(c).strip() for c in df_ag.columns]
                raw_cols = list(df_ag.columns)

                st.success(f"File loaded — **{len(df_ag)} rows**, columns: {', '.join(raw_cols)}")
                st.dataframe(df_ag.head(5), use_container_width=True, hide_index=True)

                st.markdown("**Map your columns:**")
                def _pick(label, hints, cols):
                    guess = next((c for c in cols for h in hints if h in c.lower()), cols[0])
                    return st.selectbox(label, ["(none)"] + cols, index=(cols.index(guess)+1) if guess in cols else 0, key=f"ag_map_{label}")

                none_col = "(none)"
                mc1, mc2, mc3, mc4 = st.columns(4)
                with mc1: col_session = _pick("Session / Title",  ["session","title","topic","agenda","item","description"], raw_cols)
                with mc2: col_time    = _pick("Time / Start",     ["time","start","slot","hour"], raw_cols)
                with mc3: col_speaker = _pick("Speaker",          ["speaker","presenter","name","host"], raw_cols)
                with mc4: col_track   = _pick("Track / Room",     ["track","room","stream","stage"], raw_cols)

                if col_session == none_col:
                    st.warning("Select which column contains the session title.")
                else:
                    if st.button("Save Agenda", type="primary", key="agenda_save_csv"):
                        parsed = []
                        for _, row in df_ag.iterrows():
                            parsed.append({
                                "time":    str(row[col_time]).strip()    if col_time != none_col else "",
                                "session": str(row[col_session]).strip(),
                                "speaker": str(row[col_speaker]).strip() if col_speaker != none_col else "",
                                "track":   str(row[col_track]).strip()   if col_track != none_col else "",
                            })
                        _save_agenda(parsed)
                        st.success(f"Agenda saved — {len(parsed)} sessions.")
                        st.rerun()
            else:
                st.warning("File loaded but no rows found. Is it empty?")

        # ── Add session manually ──
        with st.expander("Add a session manually"):
            mc1, mc2 = st.columns(2)
            ms_time    = mc1.text_input("Time", placeholder="e.g. 9:00 AM", key="ms_time")
            ms_session = mc2.text_input("Session title", placeholder="e.g. Opening Keynote", key="ms_session")
            ms_speaker = mc1.text_input("Speaker(s)", placeholder="e.g. Jane Smith, Acme Corp", key="ms_speaker")
            ms_track   = mc2.text_input("Track / Room", placeholder="e.g. Main Stage", key="ms_track")
            if st.button("Add Session", key="ms_add") and ms_session:
                sessions.append({
                    "time": ms_time, "session": ms_session,
                    "speaker": ms_speaker, "track": ms_track,
                })
                _save_agenda(sessions)
                st.success("Session added.")
                st.rerun()

        # ── Display agenda ──
        if not sessions:
            st.info("No agenda yet — upload a CSV or add sessions above.")
        else:
            st.markdown(f"**{len(sessions)} sessions**")

            # Group by track if tracks are present
            tracks = sorted(set(s.get("track","") for s in sessions if s.get("track")))
            if len(tracks) > 1:
                chosen_track = st.selectbox("Filter by track", ["All tracks"] + tracks, key="agenda_track_filter")
                view = sessions if chosen_track == "All tracks" else [s for s in sessions if s.get("track","") == chosen_track]
            else:
                view = sessions

            import pandas as _pd_ei
            df_agenda = _pd_ei.DataFrame(view)[["time","session","speaker","track"]]
            df_agenda.columns = ["Time","Session","Speaker","Track"]
            st.dataframe(df_agenda, use_container_width=True, hide_index=True)

            col_dl, col_clr = st.columns([3,1])
            col_dl.download_button(
                "Download Agenda CSV",
                data=df_agenda.to_csv(index=False),
                file_name=f"{_event_id}_agenda.csv",
                mime="text/csv",
            )
            if col_clr.button("Clear Agenda", key="agenda_clear"):
                _save_agenda([])
                st.rerun()

    # ── ATTENDEE LIST ─────────────────────────────────────────────────────────
    with ei_attendee_tab:
        st.subheader("Attendee List")
        st.caption(
            "Upload your registration export here. The agent reads it to build your buyer profile "
            "and understand who's in the room — so your research, scoring, and email sequences are tailored to this exact audience."
        )

        import pandas as _pd_att

        # Show current ICP summary if one exists
        current_icp = load_icp()
        if current_icp:
            m1, m2, m3 = st.columns(3)
            m1.metric("Registered Buyers", current_icp.get("buyer_count", "?"))
            m2.metric("Senior (VP+/Dir)", f"{current_icp.get('senior_buyer_pct','?')}%")
            top_cos = current_icp.get("top_companies", [])
            m3.metric("Top Companies", len(top_cos))
            if top_cos:
                st.caption("Top attending companies: " + ", ".join(top_cos[:8]))
            st.divider()

        # Drag-and-drop uploader
        att_file = st.file_uploader(
            "Drag & drop your attendee registration CSV here",
            type=["csv"],
            key="att_upload",
            help="WBR registration export. Required columns: Account, Job Title, Price List Type",
            label_visibility="collapsed",
        )

        if att_file:
            raw_att = att_file.read()
            decoded_att = raw_att.decode("utf-8-sig", errors="replace")
            df_att = _pd_att.read_csv(_io_ei.StringIO(decoded_att), on_bad_lines="skip")

            # Preview
            st.success(f"Loaded **{len(df_att)} rows** · {len(df_att.columns)} columns")

            col_prev, col_stats = st.columns(2)
            with col_prev:
                st.markdown("**Preview (first 10 rows)**")
                st.dataframe(df_att.head(10), use_container_width=True, hide_index=True)

            with col_stats:
                # Quick stats from common WBR column names
                st.markdown("**Quick stats**")
                title_col = next((c for c in df_att.columns if "title" in c.lower() or "job" in c.lower()), None)
                acct_col  = next((c for c in df_att.columns if "account" in c.lower() or "company" in c.lower()), None)
                type_col  = next((c for c in df_att.columns if "price list" in c.lower() or "type" in c.lower()), None)

                if title_col:
                    top_titles = df_att[title_col].value_counts().head(6)
                    st.caption("Top job titles:")
                    for title, cnt in top_titles.items():
                        st.caption(f"  {title} ({cnt})")

                if type_col:
                    buyers = df_att[df_att[type_col].astype(str).str.lower().str.contains("primary|buyer|delegate", na=False)]
                    vendors = df_att[df_att[type_col].astype(str).str.lower().str.contains("vendor|sponsor|exhibitor", na=False)]
                    st.caption(f"Buyers: {len(buyers)} · Sponsors: {len(vendors)}")

            st.divider()

            if st.button("Build / Rebuild Buyer Profile from this CSV", type="primary", key="att_build_icp"):
                with st.spinner("Analysing attendees..."):
                    try:
                        import event_setup as _es
                        with _tmp_ei.NamedTemporaryFile(delete=False, suffix=".csv", mode="wb") as _tmpf:
                            _tmpf.write(raw_att)
                            _tmp_path = _tmpf.name
                        new_icp = _es.build_icp_from_csv(_tmp_path)
                        storage.save_icp(new_icp, event_id=_event_id)
                        _os_ei.unlink(_tmp_path)
                        st.success(
                            f"Buyer profile updated: **{new_icp['buyer_count']} buyers** · "
                            f"**{new_icp['senior_buyer_pct']}% VP/Director+**"
                        )
                        st.rerun()
                    except Exception as _e:
                        st.error(f"Failed to build ICP: {_e}")
        else:
            if not current_icp:
                st.info(
                    "No attendee data yet. Upload your WBR registration export above to seed the buyer profile. "
                    "Once uploaded, the Research, Radar, and Outreach tabs will use the real audience data."
                )
